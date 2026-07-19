#!/usr/bin/env python3
"""
md_to_docx.py — 将 Markdown 结题报告转换为 Word 文档
===============================================
解析标题、段落、代码块、表格，生成格式化的 .docx 文件。
"""

import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


def set_cell_border(cell, **kwargs):
    """设置单元格边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = docx.oxml.OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def parse_markdown(md_text):
    """简单解析 Markdown，返回段落/表格列表。"""
    lines = md_text.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 代码块
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            blocks.append(('code', '\n'.join(code_lines)))
            i += 1
            continue

        # 表格
        if '|' in line:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            # 过滤分隔线
            data_lines = [
                ln for ln in table_lines
                if not re.fullmatch(r'\s*\|?[-:\|\s]+\|?', ln.strip())
            ]
            if data_lines:
                rows = []
                for ln in data_lines:
                    cells = [c.strip() for c in ln.split('|')]
                    cells = [c for c in cells if c or c == '']
                    # 去掉首尾空单元格（由开头/结尾的 | 产生）
                    if cells and cells[0] == '':
                        cells = cells[1:]
                    if cells and cells[-1] == '':
                        cells = cells[:-1]
                    rows.append(cells)
                blocks.append(('table', rows))
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            blocks.append(('heading', level, text))
            i += 1
            continue

        # 普通段落（合并连续非空行）
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and '|' not in lines[i] and not lines[i].strip().startswith('```'):
            para_lines.append(lines[i])
            i += 1
        blocks.append(('paragraph', '\n'.join(para_lines)))

    return blocks


def add_formatted_text(run, text):
    """处理行内 **粗体** 和 *斜体*。"""
    # 先处理粗体
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = run.font.bold = True
            run.text = part[2:-2]
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]
            run.font.italic = True
        else:
            run.text = part


def md_to_docx(md_path, docx_path):
    """转换 Markdown 文件为 Word 文档。"""
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()

    # 设置默认中文字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    blocks = parse_markdown(md_text)

    for block in blocks:
        kind = block[0]

        if kind == 'heading':
            _, level, text = block
            p = doc.add_heading(level=level)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.color.rgb = RGBColor(0, 0, 0)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = Pt(18)
                run.font.bold = True
            elif level == 2:
                run.font.size = Pt(16)
                run.font.bold = True
            elif level == 3:
                run.font.size = Pt(14)
                run.font.bold = True
            else:
                run.font.size = Pt(12)
                run.font.bold = True

        elif kind == 'paragraph':
            _, text = block
            p = doc.add_paragraph()
            # 简单处理 **粗体**
            segments = re.split(r'(\*\*[^*]+\*\*)', text)
            for seg in segments:
                run = p.add_run()
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
                if seg.startswith('**') and seg.endswith('**'):
                    run.text = seg[2:-2]
                    run.font.bold = True
                else:
                    run.text = seg
            # 首行缩进
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.paragraph_format.line_spacing = 1.5

        elif kind == 'code':
            _, code = block
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.name = 'Courier New'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.line_spacing = 1.2

        elif kind == 'table':
            _, rows = block
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            for ri, row_cells in enumerate(rows):
                row = table.rows[ri]
                for ci in range(num_cols):
                    cell = row.cells[ci]
                    if ci < len(row_cells):
                        cell.text = row_cells[ci]
                    # 设置单元格字体
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.size = Pt(10)
            doc.add_paragraph()

    # 保存
    doc.save(docx_path)
    print(f'[DOCX 生成] {docx_path}')


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base_dir, '..', '报告', '结题报告.md')
    docx_path = os.path.join(base_dir, '..', '报告', '结题报告.docx')

    if not os.path.exists(md_path):
        print(f'[错误] 找不到 Markdown 文件: {md_path}')
        sys.exit(1)

    md_to_docx(md_path, docx_path)


if __name__ == '__main__':
    main()
